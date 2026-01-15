#!/usr/bin/env python3

# Hermit - Offline AI Chatbot for Wikipedia & ZIM Files
# Copyright (C) 2026 Hermit-AI, Inc.
#
# SPDX-License-Identifier: AGPL-3.0-or-later
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published
# by the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program. If not, see <https://www.gnu.org/licenses/>.

"""
Forge - ZIM File Creator for Hermit

Create custom ZIM knowledge bases from your documents.
Supports: PDF, DOCX, TXT, Markdown, HTML, EPUB

Usage:
    forge              # Launch GUI
    forge --cli        # Interactive CLI mode
    forge /path/to/docs --output my_knowledge.zim  # CLI batch mode
"""

import os
import sys
import argparse
import hashlib
import html
from pathlib import Path
from typing import List, Dict, Optional, Generator
from datetime import datetime
import re

# Document Parsing Libraries (optional, checked at runtime)
PARSERS_AVAILABLE = {
    'pdf': False,
    'docx': False,
    'epub': False,
    'markdown': False
}

try:
    import pypdf
    PARSERS_AVAILABLE['pdf'] = True
except ImportError:
    pass

try:
    import docx
    PARSERS_AVAILABLE['docx'] = True
except ImportError:
    pass

try:
    import ebooklib
    from ebooklib import epub
    PARSERS_AVAILABLE['epub'] = True
except ImportError:
    pass

try:
    import markdown
    PARSERS_AVAILABLE['markdown'] = True
except ImportError:
    pass

# ZIM Writer
try:
    from libzim.writer import Creator, Item, StringProvider, FileProvider, Hint
    LIBZIM_AVAILABLE = True
except ImportError:
    LIBZIM_AVAILABLE = False


class Document:
    """Represents a parsed document."""
    
    def __init__(self, title: str, content: str, source_path: str, 
                 doc_type: str = "article", metadata: Dict = None):
        self.title = title
        self.content = content  # Plain text
        self.source_path = source_path
        self.doc_type = doc_type
        self.metadata = metadata or {}
        self.word_count = len(content.split())
        
        # Generate unique path for ZIM
        self.zim_path = self._generate_path()
    
    def _generate_path(self) -> str:
        """Generate a URL-safe path for this document."""
        # Clean title for URL
        safe_title = re.sub(r'[^\w\s-]', '', self.title)
        safe_title = re.sub(r'\s+', '_', safe_title).strip('_')
        
        # Add hash for uniqueness
        path_hash = hashlib.md5(self.source_path.encode()).hexdigest()[:8]
        
        return f"A/{safe_title}_{path_hash}"
    
    def to_html(self) -> str:
        """Convert document to HTML for ZIM storage."""
        # Escape content and convert newlines to paragraphs
        escaped = html.escape(self.content)
        paragraphs = escaped.split('\n\n')
        
        html_paras = '\n'.join(f'<p>{p}</p>' for p in paragraphs if p.strip())
        
        return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{html.escape(self.title)}</title>
</head>
<body>
    <h1>{html.escape(self.title)}</h1>
    <p class="metadata">Source: {html.escape(os.path.basename(self.source_path))}</p>
    <hr>
    {html_paras}
</body>
</html>
"""


class DocumentParser:
    """Parse various document formats into plain text."""
    
    @staticmethod
    def parse_file(file_path: str) -> Optional[Document]:
        """Parse a file and return a Document object."""
        path = Path(file_path)
        
        if not path.exists():
            print(f"[WARN] File not found: {file_path}")
            return None
        
        ext = path.suffix.lower()
        title = path.stem
        
        try:
            if ext == '.txt':
                content = DocumentParser._parse_txt(path)
            elif ext == '.md':
                content = DocumentParser._parse_markdown(path)
            elif ext == '.pdf':
                content = DocumentParser._parse_pdf(path)
            elif ext == '.docx':
                content = DocumentParser._parse_docx(path)
            elif ext in ['.html', '.htm']:
                content = DocumentParser._parse_html(path)
            elif ext == '.epub':
                content = DocumentParser._parse_epub(path)
            else:
                print(f"[WARN] Unsupported format: {ext}")
                return None
            
            if not content or len(content.strip()) < 10:
                print(f"[WARN] Empty or too short: {file_path}")
                return None
            
            return Document(
                title=title,
                content=content,
                source_path=str(path.absolute()),
                doc_type="article",
                metadata={"format": ext, "size": path.stat().st_size}
            )
            
        except Exception as e:
            print(f"[ERROR] Failed to parse {file_path}: {e}")
            return None
    
    @staticmethod
    def _parse_txt(path: Path) -> str:
        """Parse plain text file."""
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    @staticmethod
    def _parse_markdown(path: Path) -> str:
        """Parse Markdown file to plain text."""
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        if PARSERS_AVAILABLE['markdown']:
            # Convert MD to HTML then strip tags
            md = markdown.markdown(content)
            return re.sub(r'<[^>]+>', '', md)
        else:
            # Basic: just remove markdown syntax
            content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)
            content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
            content = re.sub(r'\*([^*]+)\*', r'\1', content)
            content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)
            return content
    
    @staticmethod
    def _parse_pdf(path: Path) -> str:
        """Parse PDF file."""
        if not PARSERS_AVAILABLE['pdf']:
            raise ImportError("pypdf not installed. Run: pip install pypdf")
        
        reader = pypdf.PdfReader(str(path))
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return '\n\n'.join(text_parts)
    
    @staticmethod
    def _parse_docx(path: Path) -> str:
        """Parse Word DOCX file."""
        if not PARSERS_AVAILABLE['docx']:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        doc = docx.Document(str(path))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return '\n\n'.join(paragraphs)
    
    @staticmethod
    def _parse_html(path: Path) -> str:
        """Parse HTML file to plain text."""
        from bs4 import BeautifulSoup
        
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        
        # Remove script and style elements
        for tag in soup(['script', 'style', 'nav', 'header', 'footer']):
            tag.decompose()
        
        return soup.get_text(separator='\n\n')
    
    @staticmethod
    def _parse_epub(path: Path) -> str:
        """Parse EPUB ebook."""
        if not PARSERS_AVAILABLE['epub']:
            raise ImportError("ebooklib not installed. Run: pip install ebooklib")
        
        from bs4 import BeautifulSoup
        
        book = epub.read_epub(str(path))
        text_parts = []
        
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), 'html.parser')
            text = soup.get_text(separator='\n')
            if text.strip():
                text_parts.append(text)
        
        return '\n\n'.join(text_parts)


class ZIMItem(Item):
    """A single item (article) in the ZIM file."""
    
    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc
        self._html = doc.to_html().encode('utf-8')
    
    def get_path(self) -> str:
        return self.doc.zim_path
    
    def get_title(self) -> str:
        return self.doc.title
    
    def get_mimetype(self) -> str:
        return "text/html"
    
    def get_contentprovider(self):
        return StringProvider(self._html.decode('utf-8'))
    
    def get_hints(self):
        return {Hint.FRONT_ARTICLE: True}


class ZIMHomePage(Item):
    """The main index page for the ZIM file."""
    
    def __init__(self, title: str, docs: List[Document]):
        super().__init__()
        self.title = title
        self.docs = docs
    
    def get_path(self) -> str:
        return "A/index"
    
    def get_title(self) -> str:
        return self.title
    
    def get_mimetype(self) -> str:
        return "text/html"
    
    def get_contentprovider(self):
        # Build index HTML
        doc_list = '\n'.join(
            f'<li><a href="{doc.zim_path}">{html.escape(doc.title)}</a> '
            f'<small>({doc.word_count} words)</small></li>'
            for doc in self.docs
        )
        
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{html.escape(self.title)}</title>
    <style>
        body {{ font-family: sans-serif; max-width: 800px; margin: 2rem auto; padding: 1rem; }}
        h1 {{ color: #333; }}
        ul {{ line-height: 1.8; }}
        a {{ color: #0066cc; text-decoration: none; }}
        a:hover {{ text-decoration: underline; }}
        small {{ color: #666; }}
    </style>
</head>
<body>
    <h1>📚 {html.escape(self.title)}</h1>
    <p>This knowledge base contains <strong>{len(self.docs)}</strong> documents.</p>
    <p>Created with <a href="https://github.com/imDelivered/Hermit-AI">Hermit Forge</a></p>
    <hr>
    <h2>Documents</h2>
    <ul>
        {doc_list}
    </ul>
</body>
</html>
"""
        return StringProvider(html_content)
    
    def get_hints(self):
        return {Hint.FRONT_ARTICLE: True}


class ZIMCreator:
    """Create ZIM files from documents."""
    
    def __init__(self, output_path: str, title: str = "My Knowledge Base"):
        if not LIBZIM_AVAILABLE:
            raise ImportError(
                "libzim not available. Install with: pip install libzim\n"
                "Or system: sudo apt install python3-libzim"
            )
        
        self.output_path = output_path
        self.title = title
        self.documents: List[Document] = []
    
    def add_document(self, doc: Document):
        """Add a parsed document."""
        self.documents.append(doc)
    
    def add_directory(self, dir_path: str, recursive: bool = True) -> int:
        """Scan a directory and add all supported files."""
        path = Path(dir_path)
        
        if not path.is_dir():
            raise ValueError(f"Not a directory: {dir_path}")
        
        extensions = {'.txt', '.md', '.pdf', '.docx', '.html', '.htm', '.epub'}
        count = 0
        
        if recursive:
            files = path.rglob('*')
        else:
            files = path.glob('*')
        
        for file_path in files:
            if file_path.suffix.lower() in extensions:
                doc = DocumentParser.parse_file(str(file_path))
                if doc:
                    self.add_document(doc)
                    count += 1
                    print(f"[OK] Added: {file_path.name}")
        
        return count
    
    def create(self, compression: str = "zstd") -> str:
        """Create the ZIM file."""
        if not self.documents:
            raise ValueError("No documents to add. Add documents first.")
        
        print(f"\n[INFO] Creating ZIM file: {self.output_path}")
        print(f"[INFO] Documents: {len(self.documents)}")
        
        # Create and configure BEFORE entering context manager
        creator = Creator(self.output_path)
        creator.config_indexing(True, "en")
        creator.set_mainpath("A/index")
        
        # Enter context manager to start writing
        with creator:
            # Add homepage
            homepage = ZIMHomePage(self.title, self.documents)
            creator.add_item(homepage)
            
            # Add all documents
            for i, doc in enumerate(self.documents):
                item = ZIMItem(doc)
                creator.add_item(item)
                
                if (i + 1) % 10 == 0:
                    print(f"[PROGRESS] Added {i + 1}/{len(self.documents)} documents")
        
        # Get file size
        size_mb = os.path.getsize(self.output_path) / (1024 * 1024)
        print(f"\n[SUCCESS] Created: {self.output_path} ({size_mb:.1f} MB)")
        
        return self.output_path


# =============================================================================
# GUI Interface
# =============================================================================

class ForgeGUI:
    """Streamlined Tkinter GUI for Forge ZIM creator."""
    
    def __init__(self, parent=None):
        try:
            import tkinter as tk
            from tkinter import ttk, filedialog, messagebox, scrolledtext
        except ImportError:
            raise RuntimeError("tkinter not available. Install: sudo apt install python3-tk")
        
        self.tk = tk
        self.ttk = ttk
        self.filedialog = filedialog
        self.messagebox = messagebox
        self.scrolledtext = scrolledtext
        
        if parent:
            self.root = tk.Toplevel(parent)
            self.root.transient(parent)
            self.root.grab_set()
        else:
            self.root = tk.Tk()
            
        self.root.title("Hermit Forge - ZIM Creator")
        self.root.geometry("600x650")
        
        # Selected folder path
        self.selected_folder: Optional[str] = None
        
        self._setup_ui()
    
    def _setup_ui(self):
        """Setup the streamlined GUI components."""
        tk = self.tk
        ttk = self.ttk
        
        # Title
        title_frame = ttk.Frame(self.root)
        title_frame.pack(fill=tk.X, padx=30, pady=20)
        
        ttk.Label(title_frame, text="🔨 Hermit Forge", font=("Helvetica", 24, "bold")).pack()
        ttk.Label(title_frame, text="Turn your documents into a searchable knowledge base", 
                  font=("Helvetica", 11), foreground="gray").pack(pady=5)
        
        # Main instruction
        instruction_frame = ttk.Frame(self.root)
        instruction_frame.pack(fill=tk.X, padx=30, pady=10)
        
        ttk.Label(instruction_frame, 
                  text="Select a folder containing your documents (PDF, TXT, MD, DOCX, HTML, EPUB)",
                  font=("Helvetica", 10), wraplength=500).pack()
        
        # Folder selection (BIG button)
        folder_frame = ttk.Frame(self.root)
        folder_frame.pack(fill=tk.X, padx=30, pady=15)
        
        self.folder_btn = ttk.Button(
            folder_frame, 
            text="📁 Select Folder", 
            command=self._select_folder,
            width=30
        )
        self.folder_btn.pack()
        
        # Selected folder display
        self.folder_label = ttk.Label(
            folder_frame, 
            text="No folder selected", 
            font=("Helvetica", 9), 
            foreground="gray"
        )
        self.folder_label.pack(pady=5)
        
        # Preview stats
        self.stats_label = ttk.Label(
            folder_frame,
            text="",
            font=("Helvetica", 9),
            foreground="#0066cc"
        )
        self.stats_label.pack()
        
        # Settings Frame (auto-populated)
        settings_frame = ttk.LabelFrame(self.root, text="ZIM Settings", padding=15)
        settings_frame.pack(fill=tk.X, padx=30, pady=10)
        
        # ZIM Title
        ttk.Label(settings_frame, text="Title:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.title_var = tk.StringVar(value="My Knowledge Base")
        ttk.Entry(settings_frame, textvariable=self.title_var, width=45).grid(row=0, column=1, padx=10, sticky=tk.W+tk.E)
        
        # Output location (auto-generated, editable)
        ttk.Label(settings_frame, text="Output:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.output_var = tk.StringVar(value="")
        output_entry = ttk.Entry(settings_frame, textvariable=self.output_var, width=45, state='readonly')
        output_entry.grid(row=1, column=1, padx=10, sticky=tk.W+tk.E)
        
        settings_frame.columnconfigure(1, weight=1)
        
        # Status/Log
        log_frame = ttk.LabelFrame(self.root, text="Progress", padding=10)
        log_frame.pack(fill=tk.BOTH, expand=True, padx=30, pady=5)
        
        self.log_text = self.scrolledtext.ScrolledText(log_frame, height=6, state=tk.DISABLED, wrap=tk.WORD)
        self.log_text.pack(fill=tk.BOTH, expand=True)
        
        # Create button
        action_frame = ttk.Frame(self.root)
        action_frame.pack(fill=tk.X, padx=30, pady=15)
        
        self.create_btn = ttk.Button(
            action_frame, 
            text="🔨 Create ZIM File", 
            command=self._create_zim,
            state=tk.DISABLED
        )
        self.create_btn.pack(side=tk.RIGHT)
    
    def _log(self, message: str):
        """Add message to log."""
        self.log_text.config(state=self.tk.NORMAL)
        self.log_text.insert(self.tk.END, message + "\n")
        self.log_text.see(self.tk.END)
        self.log_text.config(state=self.tk.DISABLED)
        self.root.update_idletasks()
    
    def _select_folder_native(self) -> Optional[str]:
        """Use native system file picker instead of Tkinter dialog."""
        import subprocess
        import shutil
        
        # Try zenity first (GNOME, XFCE, most Linux distros)
        if shutil.which('zenity'):
            try:
                result = subprocess.run(
                    ['zenity', '--file-selection', '--directory', '--title=Select Folder with Documents'],
                    capture_output=True,
                    text=True,
                    timeout=300
                )
                if result.returncode == 0:
                    return result.stdout.strip()
            except Exception:
                pass
        
        # Try kdialog (KDE Plasma)
        if shutil.which('kdialog'):
            try:
                result = subprocess.run(
                    ['kdialog', '--getexistingdirectory', os.path.expanduser('~'), '--title', 'Select Folder with Documents'],
                    capture_output=True,
                    text=True,
                    timeout=300
                )
                if result.returncode == 0:
                    return result.stdout.strip()
            except Exception:
                pass
        
        # Fallback to Tkinter dialog if native pickers aren't available
        return self.filedialog.askdirectory(title="Select Folder with Documents")
    
    def _select_folder(self):
        """Open native OS file manager to select a folder."""
        folder = self._select_folder_native()
        
        if not folder:
            return
        
        self.selected_folder = folder
        
        # Update display
        folder_name = os.path.basename(folder)
        self.folder_label.config(text=f"📁 {folder}", foreground="black")
        
        # Auto-generate ZIM title from folder name
        clean_name = folder_name.replace('_', ' ').replace('-', ' ').title()
        self.title_var.set(clean_name)
        
        # Auto-generate ZIM output path (same directory as source)
        output_path = os.path.join(folder, f"{folder_name}.zim")
        self.output_var.set(output_path)
        
        # Scan folder for supported files
        self._scan_folder(folder)
    
    def _scan_folder(self, folder: str):
        """Scan folder and show preview stats."""
        extensions = {'.txt', '.md', '.pdf', '.docx', '.html', '.htm', '.epub'}
        file_count = 0
        total_size = 0
        
        for root, _, files in os.walk(folder):
            for file in files:
                if Path(file).suffix.lower() in extensions:
                    file_count += 1
                    full_path = os.path.join(root, file)
                    try:
                        total_size += os.path.getsize(full_path)
                    except OSError:
                        pass
        
        if file_count == 0:
            self.stats_label.config(
                text="⚠ No supported documents found in this folder",
                foreground="orange"
            )
            self.create_btn.config(state=self.tk.DISABLED)
            self._log(f"No supported files found in: {folder}")
        else:
            size_mb = total_size / (1024 * 1024)
            self.stats_label.config(
                text=f"✓ Found {file_count} documents (~{size_mb:.1f} MB)",
                foreground="#00aa00"
            )
            self.create_btn.config(state=self.tk.NORMAL)
            self._log(f"Ready to process {file_count} files from: {os.path.basename(folder)}")
    
    def _create_zim(self):
        """Create the ZIM file from selected folder."""
        if not self.selected_folder:
            self.messagebox.showwarning("No Folder", "Please select a folder first.")
            return
        
        if not LIBZIM_AVAILABLE:
            self.messagebox.showerror(
                "Missing Dependency",
                "libzim is not installed.\n\n"
                "Install with:\n"
                "  sudo apt install python3-libzim\n"
                "  OR pip install libzim"
            )
            return
        
        output_path = self.output_var.get()
        title = self.title_var.get()
        
        # Check if output already exists
        if os.path.exists(output_path):
            overwrite = self.messagebox.askyesno(
                "File Exists",
                f"The file already exists:\n{output_path}\n\nOverwrite it?"
            )
            if not overwrite:
                return
        
        try:
            self.create_btn.config(state=self.tk.DISABLED)
            self._log(f"\n🔨 Creating ZIM file...")
            self._log(f"Title: {title}")
            self._log(f"Output: {output_path}")
            self._log(f"Source: {self.selected_folder}\n")
            
            # Create ZIM
            creator = ZIMCreator(output_path, title)
            
            # Manually process files to show progress in GUI
            extensions = {'.txt', '.md', '.pdf', '.docx', '.html', '.htm', '.epub'}
            processed = 0
            failed = 0
            
            folder_path = Path(self.selected_folder)
            for file_path in folder_path.rglob('*'):
                if file_path.suffix.lower() in extensions:
                    self._log(f"Processing: {file_path.name}...")
                    doc = DocumentParser.parse_file(str(file_path))
                    if doc:
                        creator.add_document(doc)
                        processed += 1
                    else:
                        failed += 1
                        self._log(f"  ⚠ Failed to parse or file too small: {file_path.name}")
            
            self._log(f"\n✓ Finished: {processed} documents added")
            if failed > 0:
                self._log(f"⚠ Details: {failed} files were skipped (likely image-based PDFs or empty files)\n")
            
            if processed == 0:
                error_msg = "No documents could be processed.\n\n"
                if not PARSERS_AVAILABLE['pdf']:
                    error_msg += "PDF support is missing. Install with:\n  pip install pypdf\n\n"
                if not PARSERS_AVAILABLE['docx']:
                    error_msg += "DOCX support is missing. Install with:\n  pip install python-docx\n\n"
                if not PARSERS_AVAILABLE['epub']:
                    error_msg += "EPUB support is missing. Install with:\n  pip install ebooklib\n\n"
                
                self.messagebox.showerror(
                    "Processing Failed",
                    error_msg
                )
                self.create_btn.config(state=self.tk.NORMAL)
                return
            
            # Create the ZIM file
            self._log("Building ZIM file (this may take a moment)...")
            zim_path = creator.create()
            
            size_mb = os.path.getsize(zim_path) / (1024 * 1024)
            self._log(f"\n✅ SUCCESS!")
            self._log(f"Created: {os.path.basename(zim_path)}")
            self._log(f"Size: {size_mb:.1f} MB")
            self._log(f"Documents: {processed}")
            
            # Auto-copy to Hermit directory
            # __file__ is /media/dekko/space/Hermit-AI-main/forge.py
            hermit_dir = os.path.dirname(os.path.abspath(__file__))
            hermit_zim_path = os.path.join(hermit_dir, os.path.basename(zim_path))
            
            try:
                import shutil
                if hermit_zim_path != zim_path:  # Only copy if different locations
                    shutil.copy2(zim_path, hermit_zim_path)
                    self._log(f"\n📥 Auto-copied to: {hermit_dir}")
                    final_path = hermit_zim_path
                else:
                    final_path = zim_path
            except Exception as e:
                self._log(f"⚠ Could not auto-copy to Hermit directory: {e}")
                final_path = zim_path
            
            # Success dialog with options
            result = self.messagebox.askyesnocancel(
                "Success!",
                f"ZIM file created successfully!\n\n"
                f"📄 {os.path.basename(final_path)}\n"
                f"📊 {size_mb:.1f} MB\n"
                f"📚 {processed} documents\n\n"
                f"✅ Ready to use with Hermit!\n\n"
                f"Yes = Open Hermit now\n"
                f"No = Open file location\n"
                f"Cancel = Just close"
            )
            
            if result is True:  # Yes - Launch Hermit
                import subprocess
                try:
                    # Try to launch hermit command
                    subprocess.Popen(['hermit'], start_new_session=True)
                    self._log("\n🚀 Launched Hermit!")
                    self.root.after(1000, self.root.destroy)  # Close Forge after 1 sec
                except Exception as e:
                    # Fallback: try running from current directory
                    try:
                        subprocess.Popen([sys.executable, 'run_chatbot.py'], 
                                       cwd=hermit_dir, start_new_session=True)
                        self._log("\n🚀 Launched Hermit!")
                        self.root.after(1000, self.root.destroy)
                    except:
                        self.messagebox.showerror("Launch Failed", 
                            f"Could not launch Hermit automatically.\n\n"
                            f"Run 'hermit' from terminal to use your new ZIM file.")
            elif result is False:  # No - Open location
                # Open file manager at the location
                import subprocess
                folder = os.path.dirname(output_path)
                try:
                    subprocess.run(['xdg-open', folder], check=False)
                except:
                    pass
            
            self.create_btn.config(state=self.tk.NORMAL)
            
        except Exception as e:
            self._log(f"\n❌ Error: {e}")
            self.messagebox.showerror("Error", str(e))
            self.create_btn.config(state=self.tk.NORMAL)
    
    def run(self):
        """Start the GUI."""
        self.root.mainloop()


# =============================================================================
# CLI Interface
# =============================================================================

def cli_main():
    """Command-line interface."""
    parser = argparse.ArgumentParser(
        description="Forge - Create ZIM knowledge bases from documents",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  forge                          # Launch GUI
  forge /path/to/docs -o kb.zim  # Convert folder to ZIM
  forge file1.pdf file2.txt -o kb.zim  # Convert specific files
        """
    )
    
    parser.add_argument("inputs", nargs="*", help="Input files or directories")
    parser.add_argument("-o", "--output", default="knowledge.zim", help="Output ZIM file path")
    parser.add_argument("-t", "--title", default="My Knowledge Base", help="Knowledge base title")
    parser.add_argument("--gui", action="store_true", help="Force GUI mode")
    parser.add_argument("--no-recursive", action="store_true", help="Don't scan subdirectories")
    
    args = parser.parse_args()
    
    # GUI mode if no inputs
    if not args.inputs or args.gui:
        print("Launching Forge GUI...")
        gui = ForgeGUI()
        gui.run()
        return
    
    # CLI mode
    if not LIBZIM_AVAILABLE:
        print("ERROR: libzim not available.")
        print("Install with: sudo apt install python3-libzim")
        sys.exit(1)
    
    creator = ZIMCreator(args.output, args.title)
    
    for input_path in args.inputs:
        path = Path(input_path)
        
        if path.is_dir():
            count = creator.add_directory(str(path), recursive=not args.no_recursive)
            print(f"Added {count} documents from {path}")
        elif path.is_file():
            doc = DocumentParser.parse_file(str(path))
            if doc:
                creator.add_document(doc)
                print(f"Added: {path.name}")
        else:
            print(f"WARNING: Not found: {input_path}")
    
    if creator.documents:
        creator.create()
    else:
        print("ERROR: No documents were added.")
        sys.exit(1)


if __name__ == "__main__":
    cli_main()